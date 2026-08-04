# NOTE: This module is no longer used. File-upload knowledge base has been removed.
# Preserved for backward compatibility with any legacy data.
"""Document parser — extract text from PDF, TXT, MD, DOCX uploads."""

from __future__ import annotations

from pathlib import Path
from typing import Optional


class DocumentParseError(Exception):
    pass


def detect_type(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext in ("pdf", "txt", "md", "markdown", "docx", "doc"):
        return "md" if ext == "markdown" else ext
    raise DocumentParseError(f"unsupported file type: .{ext}")


def parse_document(path: str | Path, file_type: Optional[str] = None) -> str:
    """Extract plain text from a document. Returns the text content."""
    path = Path(path)
    if not path.exists():
        raise DocumentParseError(f"file not found: {path}")
    ft = file_type or detect_type(path.name)

    if ft in ("txt", "md"):
        return path.read_text(encoding="utf-8", errors="replace")

    if ft == "pdf":
        # Try pdfplumber first (better text extraction, same lib as Avesia).
        # Falls back to pypdf if pdfplumber is not available or fails.
        text = _extract_pdf_pdfplumber(path)
        if text and text.strip():
            return text
        text = _extract_pdf_pypdf(path)
        if text and text.strip():
            return text
        # Last resort: OCR via VLM (gemma3:12b). Handles scanned PDFs that
        # have no text layer. Slower but works where text extraction fails.
        text = _extract_pdf_ocr(path)
        if text and text.strip():
            return text
        raise DocumentParseError(
            f"could not extract text from PDF: {path.name}. "
            "The PDF may be scanned images (no text layer) or use an unsupported encoding."
        )

    if ft in ("docx", "doc"):
        try:
            from docx import Document
        except ImportError as e:
            raise DocumentParseError("python-docx not installed") from e
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        return "\n\n".join(parts)

    raise DocumentParseError(f"unsupported file type: {ft}")


def _extract_pdf_pdfplumber(path: Path) -> str:
    """Extract text from PDF using pdfplumber (better extraction, same as Avesia).

    pdfplumber handles more PDF encodings than pypdf and preserves layout
    information (tables, columns) better. This is the primary extractor.
    """
    try:
        import pdfplumber
    except ImportError:
        return ""

    try:
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    parts.append(f"--- Page {i + 1} ---\n{txt}")
        return "\n\n".join(parts)
    except Exception:
        return ""


def _extract_pdf_pypdf(path: Path) -> str:
    """Extract text from PDF using pypdf (fallback when pdfplumber fails).

    Sometimes pdfplumber fails on certain PDFs that pypdf can handle,
    so we try both.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""

    try:
        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                parts.append(f"--- Page {i + 1} ---\n{txt}")
        return "\n\n".join(parts)
    except Exception:
        return ""


def _extract_pdf_ocr(path: Path, max_pages: int = 30) -> str:
    """Extract text from scanned PDFs using VLM (gemma3:12b via Ollama).

    This is the last-resort fallback when pdfplumber and pypdf both fail
    to extract a text layer (the PDF is scanned images). It renders each
    page as an image using pypdfium2 (PDFium) and sends them to the VLM
    for OCR.

    The VLM is the same model used for game identification (gemma3:12b),
    so no additional dependencies are needed beyond pypdfium2 (already
    installed as a pdfplumber dependency) and Ollama (already running
    for the GPCG pipeline).
    """
    import tempfile

    from gpcg.config import get_settings
    from gpcg.logging import get_logger

    log = get_logger(__name__)

    try:
        import pypdfium2 as pdfium
    except ImportError:
        log.warning("pypdfium2 not installed — cannot OCR scanned PDFs")
        return ""

    settings = get_settings()

    # Check if Ollama is available before starting
    try:
        import httpx
        resp = httpx.get(f"{settings.ollama_host}/api/tags", timeout=5)
        if resp.status_code != 200:
            log.warning("Ollama not available — cannot OCR scanned PDFs")
            return ""
    except Exception:
        log.warning("Ollama not reachable — cannot OCR scanned PDFs")
        return ""

    try:
        from gpcg.infrastructure.llm import LLMClient, LLMError
    except ImportError:
        log.warning("LLMClient not available — cannot OCR scanned PDFs")
        return ""

    vlm_model = settings.gpcg_vlm_model or "gemma3:12b"
    llm = LLMClient()

    prompt = (
        "You are an OCR engine. Extract ALL text from this page exactly as written.\n"
        "Preserve headings, paragraphs, and structure.\n"
        "Output ONLY the extracted text — no commentary, no descriptions.\n"
        "If the page has no readable text, output: [NO TEXT]"
    )

    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception as e:
        log.error(f"pypdfium2 failed to open PDF: {e}")
        return ""

    total_pages = min(len(pdf), max_pages)
    parts: list[str] = []

    try:
        for i in range(total_pages):
            page = pdf[i]
            try:
                # Render page as image at 1.5x scale for good OCR quality
                bitmap = page.render(scale=1.5)
                pil_image = bitmap.to_pil()

                # Save to temp file for the VLM
                with tempfile.NamedTemporaryFile(
                    suffix=".png", prefix=f"ocr_p{i+1}_", delete=False
                ) as tmp:
                    pil_image.save(tmp.name, "PNG")
                    tmp_path = Path(tmp.name)

                try:
                    text = llm.vision([tmp_path], prompt, model=vlm_model, temperature=0.1, max_tokens=2000)
                except LLMError as e:
                    log.warning(f"VLM OCR failed on page {i+1}: {e}")
                    text = ""
                finally:
                    tmp_path.unlink(missing_ok=True)

                # Validate VLM output: reject if the model echoed the prompt
                # instead of extracting text. This happens when the VLM
                # hallucinates the prompt text as its response.
                cleaned = (text or "").strip()
                if cleaned and "[NO TEXT]" not in cleaned[:20]:
                    # Detect prompt echo: the VLM sometimes returns the prompt
                    # itself ("You are an OCR engine. Extract ALL text...")
                    # instead of the extracted text
                    prompt_fragments = [
                        "You are an OCR engine",
                        "Extract ALL text from this page",
                        "Output ONLY the extracted text",
                        "Preserve headings, paragraphs, and structure",
                    ]
                    is_prompt_echo = any(frag in cleaned for frag in prompt_fragments)
                    if is_prompt_echo:
                        log.warning(f"VLM OCR page {i+1}: detected prompt echo, skipping")
                    else:
                        parts.append(f"--- Page {i + 1} ---\n{cleaned}")

            except Exception as e:
                log.warning(f"Failed to OCR page {i+1}: {e}")
            finally:
                page.close()

            # Log progress for long PDFs
            if (i + 1) % 5 == 0:
                log.info(f"OCR progress: {i+1}/{total_pages} pages processed")

    finally:
        pdf.close()

    result = "\n\n".join(parts)
    if result:
        log.info(f"OCR extracted text from {len(parts)}/{total_pages} pages of {path.name}")
    return result
